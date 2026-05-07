"""
Converts the internal SOP JSON structure to a styled DOCX file
matching the Varonis DAC Integration Guide template aesthetics.
"""
import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# Template color palette (from Varonis guide)
NAVY       = RGBColor(0x0F, 0x2B, 0x46)  # cover bg, H1, H3
BLUE       = RGBColor(0x1A, 0x6B, 0xB5)  # H2, links, accent
GRAY       = RGBColor(0x5A, 0x6A, 0x78)  # subtitles, metadata
DARK       = RGBColor(0x1A, 0x1A, 0x1A)  # body text
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
INFO_BG    = "E9F1F9"                     # info-box fill (hex)
COVER_BG   = "0F2B46"


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def _add_cover_page(doc: Document, sop: dict):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    cell = table.cell(0, 0)
    _set_cell_bg(cell, COVER_BG)

    # Remove borders
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Set table width
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "9360")
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    # Cell padding
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side in ["top", "left", "bottom", "right"]:
        mar = OxmlElement(f"w:{side}")
        mar.set(qn("w:w"), "720")
        mar.set(qn("w:type"), "dxa")
        tcMar.append(mar)
    tcPr.append(tcMar)

    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Clear default paragraph
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    def add_cell_para(text, size_pt, bold=False, color=WHITE, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=60):
        p = cell._tc.add_p()
        para = cell.paragraphs[-1] if cell.paragraphs else None

        from docx.text.paragraph import Paragraph
        from docx.oxml import OxmlElement as OE
        pPr = p.get_or_add_pPr()
        jc = OE("w:jc")
        align_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
        }
        jc.set(qn("w:val"), align_map.get(align, "left"))
        pPr.append(jc)

        spacing = OE("w:spacing")
        spacing.set(qn("w:before"), str(space_before))
        spacing.set(qn("w:after"), str(space_after))
        pPr.append(spacing)

        if text:
            r = OE("w:r")
            rPr = OE("w:rPr")
            # font
            rFonts = OE("w:rFonts")
            rFonts.set(qn("w:ascii"), "Calibri")
            rFonts.set(qn("w:hAnsi"), "Calibri")
            rPr.append(rFonts)
            # size
            sz = OE("w:sz")
            sz.set(qn("w:val"), str(int(size_pt * 2)))
            rPr.append(sz)
            szCs = OE("w:szCs")
            szCs.set(qn("w:val"), str(int(size_pt * 2)))
            rPr.append(szCs)
            # bold
            if bold:
                b = OE("w:b")
                rPr.append(b)
            # color
            clr = OE("w:color")
            clr.set(qn("w:val"), str(color))
            rPr.append(clr)
            r.append(rPr)
            t = OE("w:t")
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = text
            r.append(t)
            p.append(r)

    doc_type = sop.get("document_type", "Standard Operating Procedure")
    add_cell_para(doc_type.upper(), 11, bold=True, color=BLUE, space_before=0, space_after=120)
    add_cell_para(sop.get("title", "Untitled"), 26, bold=True, color=WHITE, space_before=0, space_after=160)
    add_cell_para(sop.get("subtitle", ""), 14, bold=False, color=RGBColor(0xB0, 0xC4, 0xDE), space_before=0, space_after=400)

    meta_line = "  ·  ".join(filter(None, [
        f"Version {sop.get('version', '1.0')}",
        sop.get("date", ""),
        sop.get("classification", "Internal Use Only"),
    ]))
    add_cell_para(meta_line, 10, color=RGBColor(0x9A, 0xAA, 0xB8), space_before=0, space_after=80)

    if sop.get("authors"):
        add_cell_para(f"Prepared by: {sop['authors']}", 9, color=RGBColor(0x9A, 0xAA, 0xB8), space_before=0, space_after=40)

    # Spacer
    doc.add_paragraph()


def _add_info_box(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_bg(cell, INFO_BG)

    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)

    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    run.font.name = "Calibri"
    _set_para_spacing(para, before=60, after=60)
    doc.add_paragraph()


def _apply_heading(para, level: int):
    from docx.oxml import OxmlElement as OE
    rPr_map = {
        1: {"size": 16, "color": NAVY, "bold": True},
        2: {"size": 13, "color": BLUE, "bold": True},
        3: {"size": 11, "color": NAVY, "bold": True},
    }
    cfg = rPr_map.get(level, rPr_map[3])
    for run in para.runs:
        run.font.size = Pt(cfg["size"])
        run.font.color.rgb = cfg["color"]
        run.font.bold = cfg["bold"]
        run.font.name = "Calibri"

    if level == 1:
        _set_para_spacing(para, before=280, after=120)
        # Bottom border
        pPr = para._p.get_or_add_pPr()
        pBdr = OE("w:pBdr")
        bottom = OE("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), str(NAVY))
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        _set_para_spacing(para, before=200, after=80)
    else:
        _set_para_spacing(para, before=160, after=60)


def _parse_inline(para, text: str):
    """Parse **bold**, *italic*, and `code` in a run."""
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            run = para.add_run(text[last:m.start()])
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = DARK
        if m.group(1).startswith("**"):
            run = para.add_run(m.group(2))
            run.bold = True
        elif m.group(1).startswith("*"):
            run = para.add_run(m.group(3))
            run.italic = True
        else:
            run = para.add_run(m.group(4))
            run.font.name = "Courier New"
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
        last = m.end()
    if last < len(text):
        run = para.add_run(text[last:])
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = DARK


def _add_body_para(doc: Document, text: str):
    para = doc.add_paragraph()
    _parse_inline(para, text)
    _set_para_spacing(para, before=0, after=80)
    return para


def _add_bullet(doc: Document, text: str, level: int = 0):
    para = doc.add_paragraph(style="List Paragraph")
    _parse_inline(para, text)
    _set_para_spacing(para, before=0, after=40)
    pPr = para._p.get_or_add_pPr()

    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "1")
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)
    return para


def _add_numbered(doc: Document, text: str, level: int = 0):
    para = doc.add_paragraph(style="List Paragraph")
    _parse_inline(para, text)
    _set_para_spacing(para, before=0, after=40)
    pPr = para._p.get_or_add_pPr()

    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "2")
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)
    return para


def _add_table(doc: Document, headers: list, rows: list):
    col_count = max(len(headers), max((len(r) for r in rows), default=0))
    if col_count == 0:
        return

    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, hdr in enumerate(headers[:col_count]):
        cell = hdr_row.cells[i]
        _set_cell_bg(cell, "0F2B46")
        para = cell.paragraphs[0]
        run = para.add_run(str(hdr))
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "F0F4F8" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data[:col_count]):
            cell = row.cells[c_idx]
            _set_cell_bg(cell, bg)
            para = cell.paragraphs[0]
            run = para.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.color.rgb = DARK
            run.font.name = "Calibri"

    doc.add_paragraph()


def _render_content_blocks(doc: Document, blocks: list):
    for block in blocks:
        btype = block.get("type", "paragraph")

        if btype == "paragraph":
            _add_body_para(doc, block.get("text", ""))

        elif btype == "heading":
            level = block.get("level", 2)
            style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
            para = doc.add_paragraph(block.get("text", ""), style=style_map.get(level, "Heading 2"))
            _apply_heading(para, level)

        elif btype == "bullet_list":
            for item in block.get("items", []):
                if isinstance(item, dict):
                    _add_bullet(doc, item.get("text", ""), item.get("level", 0))
                else:
                    _add_bullet(doc, str(item))

        elif btype == "numbered_list":
            for item in block.get("items", []):
                if isinstance(item, dict):
                    _add_numbered(doc, item.get("text", ""), item.get("level", 0))
                else:
                    _add_numbered(doc, str(item))

        elif btype == "table":
            _add_table(doc, block.get("headers", []), block.get("rows", []))

        elif btype == "info_box":
            _add_info_box(doc, block.get("text", ""))

        elif btype == "note":
            _add_info_box(doc, f"NOTE: {block.get('text', '')}")


def _add_numbering_xml(doc: Document):
    """Inject bullet + numbered list numbering definitions."""
    from docx.oxml import OxmlElement as OE
    part = doc.part
    try:
        numbering_part = part.numbering_part
    except Exception:
        return  # numbering already exists or not needed

    numbering = numbering_part._element

    def make_abstract(abstract_id, fmt):
        aN = OE("w:abstractNum")
        aN.set(qn("w:abstractNumId"), str(abstract_id))
        meth = OE("w:multiLevelType")
        meth.set(qn("w:val"), "hybridMultilevel")
        aN.append(meth)
        for lvl_idx in range(9):
            lvl = OE("w:lvl")
            lvl.set(qn("w:ilvl"), str(lvl_idx))
            start = OE("w:start")
            start.set(qn("w:val"), "1")
            lvl.append(start)
            numFmt = OE("w:numFmt")
            numFmt.set(qn("w:val"), fmt)
            lvl.append(numFmt)
            lvlText = OE("w:lvlText")
            lvlText.set(qn("w:val"), "•" if fmt == "bullet" else f"%{lvl_idx+1}.")
            lvl.append(lvlText)
            jc = OE("w:lvl jc") if False else OE("w:lvlJc")
            jc.set(qn("w:val"), "left")
            lvl.append(jc)
            pPr = OE("w:pPr")
            ind = OE("w:ind")
            ind.set(qn("w:left"), str(720 * (lvl_idx + 1)))
            ind.set(qn("w:hanging"), "360")
            pPr.append(ind)
            lvl.append(pPr)
            aN.append(lvl)
        return aN

    def make_num(num_id, abstract_id):
        n = OE("w:num")
        n.set(qn("w:numId"), str(num_id))
        aRef = OE("w:abstractNumId")
        aRef.set(qn("w:val"), str(abstract_id))
        n.append(aRef)
        return n

    numbering.append(make_abstract(0, "bullet"))
    numbering.append(make_abstract(1, "decimal"))
    numbering.append(make_num(1, 0))
    numbering.append(make_num(2, 1))


def generate_docx(sop: dict) -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"].font.color.rgb = DARK

    # Cover page
    _add_cover_page(doc, sop)

    # Source reference box if present
    if sop.get("source_reference"):
        _add_info_box(doc, f"SOURCE REFERENCE\n{sop['source_reference']}")

    # Sections
    for section in sop.get("sections", []):
        sec_num = section.get("number", "")
        sec_title = section.get("title", "")
        heading_text = f"{sec_num}  {sec_title}" if sec_num else sec_title

        para = doc.add_paragraph(heading_text, style="Heading 1")
        _apply_heading(para, 1)

        # Section-level content blocks
        _render_content_blocks(doc, section.get("content", []))

        # Subsections
        for sub in section.get("subsections", []):
            sub_num = sub.get("number", "")
            sub_title = sub.get("title", "")
            sub_heading = f"{sub_num}  {sub_title}" if sub_num else sub_title

            para2 = doc.add_paragraph(sub_heading, style="Heading 2")
            _apply_heading(para2, 2)

            _render_content_blocks(doc, sub.get("content", []))

            # Sub-subsections
            for subsub in sub.get("subsections", []):
                ss_num = subsub.get("number", "")
                ss_title = subsub.get("title", "")
                ss_heading = f"{ss_num}  {ss_title}" if ss_num else ss_title

                para3 = doc.add_paragraph(ss_heading, style="Heading 3")
                _apply_heading(para3, 3)

                _render_content_blocks(doc, subsub.get("content", []))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
