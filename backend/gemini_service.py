"""
Gemini service — uses the REST API directly to avoid the cffi/cryptography
dependency issue with google-generativeai on this system.
"""
import json
import os
import re
import io
import base64
import requests

GEMINI_API_BASE = "https://generativelanguage.googleapis.com/v1beta"
MODEL_TEXT  = "gemini-2.0-flash-001"
MODEL_HEAVY = "gemini-2.0-flash-001"


def _api_key() -> str:
    key = os.getenv("GEMINI_API_KEY", "")
    if not key:
        raise ValueError("GEMINI_API_KEY environment variable is not set")
    return key


def _generate(contents: list, model: str = MODEL_TEXT) -> str:
    """Call the Gemini generateContent REST endpoint with retry on 429."""
    import time
    key = _api_key()
    url = f"{GEMINI_API_BASE}/models/{model}:generateContent?key={key}"
    payload = {
        "contents": contents,
        "generationConfig": {
            "temperature": 0.3,
            "maxOutputTokens": 8192,
        },
    }
    for attempt in range(4):
        resp = requests.post(url, json=payload, timeout=120)
        if resp.status_code == 429:
            wait = 15 * (attempt + 1)  # 15s, 30s, 45s, 60s
            print(f"Rate limited — waiting {wait}s before retry {attempt + 1}/3...")
            time.sleep(wait)
            continue
        resp.raise_for_status()
        data = resp.json()
        try:
            return data["candidates"][0]["content"]["parts"][0]["text"]
        except (KeyError, IndexError) as e:
            raise RuntimeError(f"Unexpected Gemini response: {data}") from e
    raise RuntimeError("Gemini rate limit exceeded after retries. Wait a minute and try again.")


def _extract_json(text: str) -> dict:
    text = text.strip()
    fence = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", text)
    if fence:
        text = fence.group(1)
    return json.loads(text)


SOP_SCHEMA = """
{
  "title": "string",
  "subtitle": "string",
  "document_type": "string",
  "version": "string",
  "date": "string",
  "classification": "string",
  "authors": "string",
  "source_reference": "string",
  "sections": [
    {
      "number": "string",
      "title": "string",
      "content": [ ...content_blocks... ],
      "subsections": [
        {
          "number": "string",
          "title": "string",
          "content": [ ...content_blocks... ],
          "subsections": [
            { "number": "string", "title": "string", "content": [ ...content_blocks... ] }
          ]
        }
      ]
    }
  ]
}
"""

BLOCK_SCHEMA = """
Content block types:
- {"type": "paragraph", "text": "string, supports **bold** and *italic*"}
- {"type": "bullet_list", "items": ["string", ...]}
- {"type": "numbered_list", "items": ["string", ...]}
- {"type": "table", "headers": ["Col1","Col2"], "rows": [["val1","val2"],...]}
- {"type": "info_box", "text": "highlighted callout text"}
- {"type": "note", "text": "warning or note text"}
"""


def extract_and_generate_sop(raw_text: str, filename: str = "") -> dict:
    prompt = f"""You are a technical documentation specialist converting raw document content into a structured, professional Standard Operating Procedure (SOP).

Source document: {filename}

RAW CONTENT:
{raw_text[:30000]}

Generate a complete SOP JSON following EXACTLY this schema:
{SOP_SCHEMA}

{BLOCK_SCHEMA}

Rules:
1. Preserve ALL technical details from the source document.
2. Organize into logical numbered sections (1, 2, 3) and subsections (1.1, 1.2).
3. Convert prose instructions into bullet or numbered lists where appropriate.
4. Use info_box for important notes, warnings, or callouts.
5. Use tables for comparisons, requirements, or structured data.
6. Output ONLY valid JSON — no markdown fences, no prose.
"""
    text = _generate([{"role": "user", "parts": [{"text": prompt}]}])
    return _extract_json(text)


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text_from_pdf_bytes(file_bytes: bytes, filename: str) -> str:
    """Upload the PDF to Gemini File API then extract text."""
    key = _api_key()

    # 1. Upload file
    upload_url = f"https://generativelanguage.googleapis.com/upload/v1beta/files?key={key}"
    headers = {
        "X-Goog-Upload-Command": "start, upload, finalize",
        "X-Goog-Upload-Header-Content-Type": "application/pdf",
        "X-Goog-Upload-Header-Content-Length": str(len(file_bytes)),
        "Content-Type": "application/pdf",
    }
    upload_resp = requests.post(upload_url, headers=headers, data=file_bytes, timeout=120)
    upload_resp.raise_for_status()
    file_uri = upload_resp.json()["file"]["uri"]

    # 2. Extract text
    prompt = "Extract ALL text content from this PDF. Preserve headings, lists, tables, and paragraphs. Output plain text only."
    contents = [{"role": "user", "parts": [
        {"file_data": {"mime_type": "application/pdf", "file_uri": file_uri}},
        {"text": prompt},
    ]}]
    return _generate(contents)


def chat_refine(sop: dict, message: str, history: list) -> tuple[dict, str]:
    history_text = ""
    for h in history[-6:]:
        history_text += f"{h.get('role','user').upper()}: {h.get('content','')}\n"

    prompt = f"""You are editing a Standard Operating Procedure document.

CURRENT SOP JSON:
{json.dumps(sop, indent=2)[:15000]}

CONVERSATION HISTORY:
{history_text}

USER REQUEST: {message}

Apply the change. Return a JSON object with exactly two keys:
1. "sop": the complete updated SOP JSON (same schema as input)
2. "reply": a brief (1-2 sentence) confirmation of what you changed

Output ONLY valid JSON — no markdown fences.
"""
    text = _generate([{"role": "user", "parts": [{"text": prompt}]}])
    result = _extract_json(text)
    return result.get("sop", sop), result.get("reply", "Done! Document updated as requested.")


def generate_blank_sop(title: str = "New SOP") -> dict:
    return {
        "title": title,
        "subtitle": "Enter subtitle here",
        "document_type": "Standard Operating Procedure",
        "version": "1.0",
        "date": "2026",
        "classification": "Internal Use Only",
        "authors": "",
        "source_reference": "",
        "sections": [
            {
                "number": "1",
                "title": "Overview",
                "content": [{"type": "paragraph", "text": "Describe the purpose of this SOP here."}],
                "subsections": [
                    {
                        "number": "1.1",
                        "title": "Scope",
                        "content": [{"type": "paragraph", "text": "Define the scope of this procedure."}],
                        "subsections": [],
                    }
                ],
            },
            {
                "number": "2",
                "title": "Prerequisites",
                "content": [{"type": "bullet_list", "items": ["Requirement 1", "Requirement 2"]}],
                "subsections": [],
            },
            {
                "number": "3",
                "title": "Procedure",
                "content": [{"type": "paragraph", "text": "Describe the step-by-step procedure."}],
                "subsections": [
                    {
                        "number": "3.1",
                        "title": "Step 1",
                        "content": [{"type": "numbered_list", "items": ["Do this first", "Then do this", "Finally do this"]}],
                        "subsections": [],
                    }
                ],
            },
        ],
    }
